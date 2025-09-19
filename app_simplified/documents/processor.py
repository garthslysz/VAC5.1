"""
Document processing for VAC assessments
"""

import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
import asyncio
from datetime import datetime
import uuid

from fastapi import UploadFile
import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Processes uploaded documents for VAC assessments
    Handles PDF, DOCX, and text files
    """
    
    def __init__(self):
        self.processed_files = {}  # In-memory storage for demo
        self.case_files = {}  # Maps case_id to list of file_ids
        self.upload_dir = Path("data/uploads")
        self.upload_dir.mkdir(exist_ok=True)
    
    async def process_file(
        self,
        file: UploadFile,
        case_id: Optional[str] = None,
        user_id: str = "anonymous"
    ) -> Dict[str, Any]:
        """
        Process uploaded file and extract text content
        
        Args:
            file: Uploaded file
            case_id: Associated case ID
            user_id: User who uploaded the file
            
        Returns:
            Processing result with extracted content
        """
        try:
            file_id = str(uuid.uuid4())
            filename = file.filename or f"upload_{file_id}"
            
            # Read file content
            content = await file.read()
            
            # Save file to disk for persistence
            file_path = self.upload_dir / f"{file_id}_{filename}"
            with open(file_path, "wb") as f:
                f.write(content)
            
            # Reset file pointer for potential re-reading
            await file.seek(0)
            
            # Determine file type and extract text
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == ".pdf":
                extracted_text = await self._extract_pdf_text(content)
                file_type = "PDF"
            elif file_extension in [".docx", ".doc"]:
                extracted_text = await self._extract_docx_text(content)
                file_type = "Word Document"
            elif file_extension == ".txt":
                extracted_text = content.decode('utf-8', errors='ignore')
                file_type = "Text Document"
            elif file_extension == ".json":
                extracted_text = content.decode('utf-8', errors='ignore')
                file_type = "JSON Document"
            else:
                # Try to decode as text
                try:
                    extracted_text = content.decode('utf-8', errors='ignore')
                    file_type = "Text Document"
                except:
                    raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Analyze content for medical conditions
            medical_analysis = await self._analyze_medical_content(extracted_text)
            
            # Store processed file info
            file_info = {
                "file_id": file_id,
                "filename": filename,
                "file_type": file_type,
                "file_size": len(content),
                "file_path": str(file_path),
                "extracted_text": extracted_text,
                "text_length": len(extracted_text),
                "medical_analysis": medical_analysis,
                "case_id": case_id,
                "user_id": user_id,
                "processed_at": datetime.now().isoformat(),
                "status": "processed"
            }
            
            self.processed_files[file_id] = file_info
            
            # Associate with case if provided
            if case_id:
                if case_id not in self.case_files:
                    self.case_files[case_id] = []
                self.case_files[case_id].append(file_id)
            
            logger.info(f"Successfully processed file: {filename} ({file_type})")
            
            return {
                "file_id": file_id,
                "filename": filename,
                "file_type": file_type,
                "status": "processed",
                "text_length": len(extracted_text),
                "conditions_detected": medical_analysis.get("conditions_detected", []),
                "processing_time": "immediate"
            }
            
        except Exception as e:
            logger.error(f"File processing error for {filename}: {e}")
            
            # Store error info
            error_file_info = {
                "file_id": str(uuid.uuid4()),
                "filename": filename,
                "status": "failed",
                "error": str(e),
                "processed_at": datetime.now().isoformat()
            }
            
            return {
                "filename": filename,
                "status": "failed",
                "error": str(e)
            }
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
            raise ValueError(f"Could not extract text from PDF: {str(e)}")
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            import io
            doc = DocxDocument(io.BytesIO(content))
            
            text_content = []
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX text extraction error: {e}")
            raise ValueError(f"Could not extract text from Word document: {str(e)}")
    
    async def _analyze_medical_content(self, text: str) -> Dict[str, Any]:
        """
        Analyze extracted text for medical conditions and relevant information
        Simple keyword-based analysis for demo
        """
        if not text:
            return {"conditions_detected": [], "analysis_confidence": "low"}
        
        text_lower = text.lower()
        
        # Common medical conditions that might appear in VAC assessments
        condition_keywords = {
            "ptsd": ["ptsd", "post-traumatic stress", "post traumatic stress"],
            "depression": ["depression", "depressive", "major depressive"],
            "anxiety": ["anxiety", "anxiety disorder", "generalized anxiety"],
            "back_pain": ["back pain", "lower back", "lumbar", "spine"],
            "knee_injury": ["knee", "knee injury", "knee pain", "patella"],
            "shoulder_injury": ["shoulder", "shoulder injury", "rotator cuff"],
            "tinnitus": ["tinnitus", "ringing in ears", "ear ringing"],
            "hearing_loss": ["hearing loss", "hearing impairment", "deaf"],
            "arthritis": ["arthritis", "osteoarthritis", "rheumatoid"],
            "migraine": ["migraine", "headache", "chronic headaches"],
            "sleep_disorder": ["insomnia", "sleep disorder", "sleep apnea"],
            "traumatic_brain_injury": ["tbi", "traumatic brain injury", "head injury"],
        }
        
        detected_conditions = []
        for condition, keywords in condition_keywords.items():
            for keyword in keywords:
                if keyword in text_lower:
                    detected_conditions.append({
                        "condition": condition,
                        "keyword_matched": keyword,
                        "confidence": "medium"
                    })
                    break  # Don't add same condition multiple times
        
        # Look for symptom indicators
        symptom_keywords = [
            "pain", "difficulty", "unable", "impairment", "limitation",
            "chronic", "severe", "moderate", "mild", "symptoms", "diagnosed"
        ]
        
        symptoms_found = [kw for kw in symptom_keywords if kw in text_lower]
        
        # Assess document type
        doc_type = "unknown"
        if any(term in text_lower for term in ["medical report", "doctor", "physician", "diagnosis"]):
            doc_type = "medical_report"
        elif any(term in text_lower for term in ["assessment", "evaluation", "rating"]):
            doc_type = "assessment_document"
        elif any(term in text_lower for term in ["service", "military", "deployment"]):
            doc_type = "service_record"
        
        return {
            "conditions_detected": detected_conditions,
            "symptoms_found": symptoms_found,
            "document_type": doc_type,
            "analysis_confidence": "medium" if detected_conditions else "low",
            "text_sample": text[:200] + "..." if len(text) > 200 else text,
            "medical_indicators": len(detected_conditions) + len(symptoms_found)
        }
    
    async def get_case_files(self, case_id: str) -> List[Dict[str, Any]]:
        """Get all files associated with a case"""
        try:
            file_ids = self.case_files.get(case_id, [])
            
            case_files = []
            for file_id in file_ids:
                file_info = await self.get_file_content(file_id)
                if file_info:
                    # Return summary info (not full extracted text)
                    case_files.append({
                        "file_id": file_id,
                        "filename": file_info["filename"],
                        "file_type": file_info["file_type"],
                        "status": file_info["status"],
                        "processed_at": file_info["processed_at"],
                        "text_length": file_info.get("text_length", 0),
                        "conditions_detected": file_info.get("medical_analysis", {}).get("conditions_detected", []),
                        "document_type": file_info.get("medical_analysis", {}).get("document_type", "unknown")
                    })
            
            return case_files
            
        except Exception as e:
            logger.error(f"Error getting case files for {case_id}: {e}")
            return []
    
    async def get_file_content(self, file_id: str) -> Optional[Dict[str, Any]]:
        """Get full file content including extracted text"""
        file_info = self.processed_files.get(file_id)
        if file_info:
            return file_info
        
        # Check if file exists on disk
        for file_path in self.upload_dir.glob(f"{file_id}_*"):
            if file_path.exists():
                # Re-extract content from disk
                with open(file_path, "rb") as f:
                    content = f.read()
                
                filename = file_path.name.replace(f"{file_id}_", "")
                file_type = self._get_file_type(filename)
                
                # Re-extract text
                extracted_text = await self._extract_text_from_content(content, filename)
                medical_analysis = await self._analyze_medical_content(extracted_text)
                
                file_info = {
                    "file_id": file_id,
                    "filename": filename,
                    "file_type": file_type,
                    "file_size": len(content),
                    "file_path": str(file_path),
                    "extracted_text": extracted_text,
                    "text_length": len(extracted_text),
                    "medical_analysis": medical_analysis,
                    "processed_at": datetime.now().isoformat(),
                    "status": "processed"
                }
                
                # Cache in memory
                self.processed_files[file_id] = file_info
                return file_info
        
        return None
    
    def _get_file_type(self, filename: str) -> str:
        """Determine file type from filename"""
        file_extension = Path(filename).suffix.lower()
        if file_extension == ".pdf":
            return "PDF"
        elif file_extension in [".docx", ".doc"]:
            return "Word Document"
        elif file_extension == ".txt":
            return "Text Document"
        elif file_extension == ".json":
            return "JSON Document"
        else:
            return "Unknown"
    
    async def _extract_text_from_content(self, content: bytes, filename: str) -> str:
        """Extract text from file content"""
        file_extension = Path(filename).suffix.lower()
        
        if file_extension == ".pdf":
            return await self._extract_pdf_text(content)
        elif file_extension in [".docx", ".doc"]:
            return await self._extract_docx_text(content)
        elif file_extension == ".txt":
            return content.decode('utf-8', errors='ignore')
        elif file_extension == ".json":
            return content.decode('utf-8', errors='ignore')
        else:
            try:
                return content.decode('utf-8', errors='ignore')
            except:
                return "Unable to extract text from file"
    
    async def delete_file(self, file_id: str) -> bool:
        """Delete a processed file"""
        try:
            if file_id in self.processed_files:
                file_info = self.processed_files[file_id]
                file_path = file_info.get("file_path")
                if file_path and Path(file_path).exists():
                    Path(file_path).unlink()
                
                # Remove from case associations
                for case_id, file_list in self.case_files.items():
                    if file_id in file_list:
                        file_list.remove(file_id)
                
                # Remove file info
                del self.processed_files[file_id]
                return True
            
            # Check if file exists on disk
            for file_path in self.upload_dir.glob(f"{file_id}_*"):
                if file_path.exists():
                    file_path.unlink()
                    return True
            
            return False
            
        except Exception as e:
            logger.error(f"Error deleting file {file_id}: {e}")
            return False
    
    def get_stats(self) -> Dict[str, Any]:
        """Get processing statistics"""
        total_files = len(self.processed_files)
        
        # Count files on disk
        disk_files = list(self.upload_dir.glob("*"))
        total_disk_files = len(disk_files)
        
        successful = len([f for f in self.processed_files.values() if f.get("status") == "processed"])
        failed = total_files - successful
        
        file_types = {}
        for file_info in self.processed_files.values():
            file_type = file_info.get("file_type", "unknown")
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            "total_files": total_files,
            "total_disk_files": total_disk_files,
            "successful": successful,
            "failed": failed,
            "file_types": file_types,
            "cases_with_files": len(self.case_files)
        }

# Global instance for application use
document_processor = DocumentProcessor()
